from docx import Document
from docx.shared import Pt
import os

def create_cv(name,email,phone,title1,title2,
              location,profile,education1,education2,
              certificate1,certificate2,position1,institution1,
              duration_beg1,duration_end1,workdone1,position2,
              institution2,duration_beg2,duration_end2,workdone2,
              language1,language2,skill1,skill2):
    
    doc=Document()

    # set font
    style=doc.styles['Normal']
    font=style.font
    font.name='Times New Roman'
    font.size=Pt(12)

    # Add Header
    doc.add_heading(name,0)
    doc.add_paragraph(f" {title1} | {title2}")
    doc.add_paragraph(f" 📍 {location}  📧 {email} 📞 {phone}")


    # Placeholders sections
    doc.add_heading('PROFILE',level=1)
    doc.add_paragraph(profile)

# education
    doc.add_heading("EDUCATION", level=1)
    doc.add_paragraph(f"{education1} \n {education2} ")

# certificate
    doc.add_heading("CERTIFICATE",level=1)
    doc.add_paragraph(f"{certificate1} \n {certificate2}")


    # experince
    doc.add_heading("EXPERIENCE",level=1)

    doc.add_paragraph(f"{position1}\n{institution1} \n {duration_beg1} - {duration_end1}")
    doc.add_paragraph(workdone1)

    doc.add_paragraph(f"{position2}\n{institution2} \n {duration_beg2 } - {duration_end2}")
    doc.add_paragraph(workdone2)

# skills
    doc.add_heading("SKILLS",level=1)

    doc.add_paragraph(f"{language1} \n {language2}")
    doc.add_paragraph(f"{skill1} \n {skill2}")


    save_to="downloads"
    os.makedirs(save_to,exist_ok=True)

    filename=os.path.join(save_to,f"{name.replace('','_')}_CV.docx")
    doc.save(filename)
    return filename